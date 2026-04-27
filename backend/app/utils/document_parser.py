from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.exceptions import DocumentEmptyError, DocumentParseFailedError
from app.schemas.domain import FileInfo, ProviderName


class DocumentParser:
    def __init__(self, provider_name: ProviderName) -> None:
        self.provider_name = provider_name.value

    def extract_text(self, file_info: FileInfo) -> str:
        path = Path(file_info.path)
        try:
            if file_info.extension == ".docx":
                text = self._extract_docx(path)
            elif file_info.extension == ".pdf":
                text = self._extract_pdf(path)
            else:
                raise DocumentParseFailedError(
                    self.provider_name,
                    detail={
                        "filename": file_info.original_filename,
                        "path": str(path),
                        "extension": file_info.extension,
                    },
                )
        except DocumentEmptyError:
            raise
        except Exception as exc:
            raise DocumentParseFailedError(
                self.provider_name,
                detail={
                    "filename": file_info.original_filename,
                    "path": str(path),
                    "extension": file_info.extension,
                    "reason": str(exc),
                },
            ) from exc

        normalized = text.strip()
        if not normalized:
            raise DocumentEmptyError(
                self.provider_name,
                detail={
                    "filename": file_info.original_filename,
                    "path": str(path),
                    "extension": file_info.extension,
                },
            )
        return normalized

    @staticmethod
    def _extract_docx(path: Path) -> str:
        document = DocxDocument(path)
        chunks: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    chunks.append(row_text)

        return "\n\n".join(chunks)

    @staticmethod
    def _extract_pdf(path: Path) -> str:
        reader = PdfReader(path)
        chunks: list[str] = []

        for page in reader.pages:
            page_text = (page.extract_text() or "").strip()
            if page_text:
                chunks.append(page_text)

        return "\n\n".join(chunks)
